"""
MONOGRAFIAS REVISOR EN MASA
"""

#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
import pdfplumber

# 1️⃣ Cargar API Key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# 2️⃣ Función para leer DOCX
def leer_docx(ruta):
    doc = Document(ruta)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

# 3️⃣ Función para leer PDF
def leer_pdf(ruta):
    texto = []
    with pdfplumber.open(ruta) as pdf:
        for pagina in pdf.pages:
            contenido = pagina.extract_text()
            if contenido:
                texto.append(contenido.strip())
    return "\n".join(texto)

# 4️⃣ Prompt base (tono académico)
prompt_base_tono = """
Actúa como un revisor académico experto en metodología de investigación cuantitativa.
Tu tarea es hacer una revisión detallada, minuciosa y objetiva de la monografía proporcionada.

Instrucciones importantes sobre el tono y la forma:
- No suavices ni endulces las observaciones.
- No digas que algo está "bien" si hay errores, incoherencias o falta de claridad.
- No hagas las observaciones más duras ni más suaves: descríbelas tal cual son, con objetividad y rigor académico.
- No intentes motivar ni dar mensajes positivos. El objetivo es que la retroalimentación sea precisa, clara y útil para corregir el trabajo.
- Usa un lenguaje formal, académico y directo.

Devuelve únicamente las observaciones con tono académico neutral y objetivo. No incluyas frases de ánimo ni de felicitación.
"""

# 5️⃣ Prompts de agentes con la frase añadida en cada uno
frase_control = """Actúa como un revisor académico experto en metodología de investigación cuantitativa.
No propones ejemplos de mejoras ni sugerencias.
Te limitas únicamente a hacer observaciones, no propones mejoras ni haces sugerencias.
"""

prompts_agentes = {
    "TITULO": f"""{frase_control}
Evalúa en el documento el TÍTULO de la monografía considerando:
- Claridad y precisión en el tema central.
- Identificación de variables independientes y dependientes.
- Delimitación geográfica y temporal.
- Coherencia con el objetivo general y el planteamiento del problema.
- Concisión (15-25 palabras).
- Ejemplo de estructura correcta: “Relación entre la frecuencia de compra y la rotación de inventario en Totto Potosí, gestión 2024”.
Entrega observaciones detalladas.
""",
    "INTRODUCCIÓN": f"""{frase_control}
Evalúa en el documento la INTRODUCCIÓN considerando el TÍTULO:
- Presenta claramente el tema central.
- Explica el contexto general y la importancia del estudio.
- Introduce las variables, ámbito geográfico y temporal.
- Conecta con el problema de investigación.
- Justificación inicial coherente.
- Relación con CONTEXTO, OBJETIVOS y PLANTEAMIENTO DEL PROBLEMA.
Entrega observaciones detalladas.
""",
    "CONTEXTO Y JUSTIFICACIÓN": f"""{frase_control}
Evalúa en el documento CONTEXTO Y JUSTIFICACIÓN considerando TÍTULO e INTRODUCCIÓN:
- Descripción clara de la empresa, sector o mercado.
- Relevancia para la gestión comercial y uso de Business Intelligence.
- Justificación social donde se explica el aporte de esta monografía a la sociedad en general.
- Justificación económica donde se explica el aporte en términos de mejora económica de esta monografía a la sociedad en general.
- Justificación comercial donde se explica el aporte de mejora comercial de esta monografía a la empresa y la sociedad en general.
- Inclusión de referencias o datos de apoyo.
- Coherencia con INTRODUCCIÓN, PLANTEAMIENTO DEL PROBLEMA y OBJETIVOS.
Entrega observaciones detalladas.
""",
    "PLANTEAMIENTO DEL PROBLEMA": f"""{frase_control}
Evalúa en el documento PLANTEAMIENTO DEL PROBLEMA considerando TÍTULO, INTRODUCCIÓN y CONTEXTO:
- Definición clara y concreta del problema.
- Redacción como pregunta de investigación.
- Identificación de causas y consecuencias.
- Variables independientes y dependientes.
- Delimitación temporal y geográfica.
- Coherencia con OBJETIVOS y DIAGNÓSTICO.
- Verifica si la pregunta problemática está correctamente formulada.
Entrega observaciones detalladas.
""",
    "OBJETIVOS": f"""{frase_control}
Evalúa en el documento OBJETIVO GENERAL y OBJETIVOS ESPECÍFICOS:
- Claridad y redacción con verbos en infinitivo.
- Relación directa con variables y problema.
- Delimitación temporal y geográfica.
- Secuencia lógica de objetivos específicos para alcanzar el general.
- Coherencia con PLANTEAMIENTO DEL PROBLEMA, ALCANCE y METODOLOGÍA.
Entrega observaciones detalladas.
""",
    "ALCANCE Y LÍMITES": f"""{frase_control}
Evalúa en el documento ALCANCE Y LÍMITES considerando OBJETIVOS y PROBLEMA:
- Definición de población, variables, lugar y tiempo.
- Coherencia con objetivos y metodología.
- Diferenciación clara entre alcance (lo que cubre) y límites (lo que excluye).
- Razonabilidad de los límites definidos.
Entrega observaciones detalladas.
""",
    "MARCO TEÓRICO": f"""{frase_control}
Evalúa en el documento MARCO TEÓRICO REFERENCIAL considerando CONTEXTO, PROBLEMA y OBJETIVOS:
- 6.1. Business Intelligence y Análisis de Datos: Conceptos clave, autores y su relación con el problema.
- 6.2. Funcionalidades de Power BI para Análisis: Profundidad técnica y ejemplos aplicados.
- 6.3. Indicadores Clave de Desempeño (KPIs): Definición y explicación de cada uno de los KPIs utilizados en el análisis y los gráficos.
- 6.4. Modelos de Pronóstico en Excel: Explicación teórica y vínculo con objetivos.
- Uso de citas y formato APA.
- Coherencia general con el trabajo.
Entrega observaciones detalladas para cada subpunto.
""",
    "DIAGNÓSTICO": f"""{frase_control}
Evalúa en el documento DIAGNÓSTICO DE LA SITUACIÓN ACTUAL considerando CONTEXTO, PROBLEMA y OBJETIVOS:
- 7.1. Descripción de la Empresa y Proceso de Ventas: Datos reales, tablas y gráficos.
- 7.2. Diccionario de datos: Explicación y descripción clara de cada una de las variables numéricas y categóricas utilizadas en el análisis.
- 7.3. Estructura y Limpieza de la Base de Datos: Descripción clara de la preparación de datos.
- 7.4. Calidad y Consistencia: Evaluación de la integridad de los datos.
- Relación del diagnóstico con el problema.
- Coherencia con la METODOLOGÍA y ANÁLISIS DE DATOS.
Entrega observaciones detalladas para cada subpunto.
""",
    "METODOLOGÍA": f"""{frase_control}
Evalúa en el documento la METODOLOGÍA considerando OBJETIVOS y ALCANCE:
- 8.1. Extracción y limpieza de Datos con Power Query: Procedimientos claros y replicables.
- 8.2. Modelado en Power BI: Estructura de datos, relaciones y presentación de gráfico de relaciones.
- Claridad en pasos, herramientas utilizadas y replicabilidad.
- Coherencia con el DIAGNÓSTICO y ANÁLISIS DE DATOS.
Entrega observaciones detalladas para cada subpunto.
""",
    "ANÁLISIS DE DATOS": f"""{frase_control}
Evalúa en el documento el ANÁLISIS DE DATOS considerando METODOLOGÍA y OBJETIVOS:
- 9.1. Análisis Descriptivo: Explicación clara de los puntos que seran aboradados acontinuación en el subtitulo.
- 9.2. Análisis Univariado: Presentación con gráficos, interpretación y explicación clara de variables individuales numéricas y categóricas.
- 9.3. Análisis Bivariado y Correlaciones: Presentación con gráficos, interpretación y explicación clara de relaciones entre variables.
- 9.4. Análisis de Indicadores Comerciales: Presentación con gráficos, interpretación y explicación clara de la relación entre variables y su relevancia para la gestión.
- 9.5. Segmentación Avanzada: Uso de técnicas de agrupación y segmentación.
- Conexión de los resultados con los OBJETIVOS y PROBLEMA.
Entrega observaciones detalladas para cada subpunto.
""",
    "MODELOS Y PRONÓSTICOS": f"""{frase_control}
Evalúa en el documento GENERACIÓN DE MODELOS Y PRONÓSTICOS considerando ANÁLISIS DE DATOS y OBJETIVOS:
- 10.1. Regresión Lineal con indicadores comerciales: Correcta aplicación y explicación.
- 10.2. Pronóstico con Función TENDENCIA: Pertinencia y resultados.
- Interpretación de resultados y conexión con los OBJETIVOS.
Entrega observaciones detalladas para cada subpunto.
""",
    "CONCLUSIONES": f"""{frase_control}
Evalúa en el documento CONCLUSIONES Y RECOMENDACIONES considerando todo el trabajo:
- Responden al objetivo general y específicos.
- Resumen claro de los hallazgos principales.
- Recomendaciones prácticas y aplicables.
- Coherencia global con TÍTULO, PROBLEMA, ANÁLISIS y MODELOS.
Entrega observaciones detalladas.
""",

    "CALIFICACIÓN FINAL": f"""{frase_control}
Evalúa en el documento considerando todo el trabajo y todas las observaciones:
- Otorga una calificacón que se encuentre entre 0% y 100%, donde 0% es PESIMO  y 100% es EXCELENTE
Entrega las observaciones mas destacadas que jsutifiquen la calificacion asignada.
"""
}

# 6️⃣ Prompt de coherencia global
prompt_coherencia_global = f"""{prompt_base_tono}

Evalúa la COHERENCIA GLOBAL considerando:
- Consistencia entre TÍTULO, INTRODUCCIÓN, CONTEXTO, PROBLEMA y OBJETIVOS.
- Correspondencia de variables independientes y dependientes a lo largo del documento.
- Coherencia en delimitación geográfica y temporal.
- Ausencia de contradicciones entre secciones.
- Fluidez lógica entre capítulos.

Entrega observaciones detalladas sobre la coherencia general del documento.
"""

# 7️⃣ Función para ejecutar agente
def ejecutar_agente(seccion, prompt, texto, contexto_previos):
    contenido = prompt_base_tono + "\n\n" + prompt + "\n\n"
    if contexto_previos:
        contenido += "ANÁLISIS DE AGENTES PREVIOS:\n" + "\n".join(contexto_previos) + "\n\n"
    contenido += "TEXTO DE LA MONOGRAFÍA:\n" + texto

    respuesta = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"Eres un revisor experto en {seccion}."},
            {"role": "user", "content": contenido}
        ],
        temperature=0.2,
        max_tokens=5000
    )
    return respuesta.choices[0].message.content

# 8️⃣ Analizar monografía
def analizar_monografia_agentes(ruta_archivo):
    ext = ruta_archivo.lower().split(".")[-1]
    if ext == "docx":
        texto = leer_docx(ruta_archivo)
    elif ext == "pdf":
        texto = leer_pdf(ruta_archivo)
    else:
        print(f"⏩ Archivo no soportado: {ruta_archivo}")
        return

    nombre = os.path.splitext(os.path.basename(ruta_archivo))[0]
    carpeta_salida = os.path.dirname(ruta_archivo)
    ruta_out = os.path.join(carpeta_salida, f"INFORME_REVISION_{nombre}.docx")

    doc = Document()
    doc.add_heading('Informe de Revisión de Monografía', 0)

    contexto_previos = []

    # Ejecutar agentes sección por sección
    for seccion, prompt in prompts_agentes.items():
        doc.add_heading(seccion, level=1)
        resultado = ejecutar_agente(seccion, prompt, texto, contexto_previos)
        contexto_previos.append(f"[{seccion}]\n{resultado}")
        for linea in resultado.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())

    # Agente final de coherencia global
    doc.add_heading("COHERENCIA GLOBAL", level=1)
    resultado_global = ejecutar_agente("Coherencia Global", prompt_coherencia_global, texto, contexto_previos)
    for linea in resultado_global.split("\n"):
        if linea.strip():
            doc.add_paragraph(linea.strip())

    doc.save(ruta_out)
    print(f"✅ Informe generado: {ruta_out}")

# 9️⃣ Procesar todas las carpetas
def procesar_todas_monografias(carpeta_base):
    for carpeta in os.listdir(carpeta_base):
        ruta_carpeta = os.path.join(carpeta_base, carpeta)
        if os.path.isdir(ruta_carpeta):
            archivos_validos = [f for f in os.listdir(ruta_carpeta) if f.lower().endswith((".docx", ".pdf"))]
            if archivos_validos:
                for archivo in archivos_validos:
                    ruta_archivo = os.path.join(ruta_carpeta, archivo)
                    analizar_monografia_agentes(ruta_archivo)
            else:
                print(f"⚠️ No hay archivos DOCX o PDF en {ruta_carpeta}")

# 🔟 Ejecutar
if __name__ == "__main__":
    carpeta_monografias = r"C:\Users\HP\Downloads\MONOGRAFIAS"
    procesar_todas_monografias(carpeta_monografias)


# In[ ]:




